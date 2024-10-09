import sys
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.firefox.service import Service
from selenium.common.exceptions import NoSuchElementException
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.oxml.ns import qn
import time
import traceback


def parser2(string_to_put_in_search: str, date_from: str, date_up: str, flag: str):
    def add_hyperlink(paragraph, text, url):
        # Получаем идентификатор ссылки (relationship id)
        rel_id = paragraph.part.relate_to(url, 'hyperlink', is_external=True)

        # Создаем элемент гиперссылки
        hyperlink = OxmlElement('w:hyperlink', {qn('r:id'): rel_id})

        # Создаем элемент 'run' и 'text'
        run = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = text
        run.append(t)

        # Создаем свойства для текста (например, подчеркивание и цвет)
        run_properties = OxmlElement('w:rPr')

        # Подчеркивание
        underline = OxmlElement('w:u')
        underline.set(qn('w:val'), 'single')
        run_properties.append(underline)

        # Цвет шрифта
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')  # Синий цвет
        run_properties.append(color)

        # Добавляем свойства к run
        run.append(run_properties)

        # Добавляем run в гиперссылку
        hyperlink.append(run)

        # Добавляем гиперссылку в параграф
        paragraph._element.append(hyperlink)

    def parse_SC(search_string: str, date_from: str, date_up: str) -> list:
        # Настройка Selenium
        firefox_options = webdriver.FirefoxOptions()
        firefox_options.add_argument('--headless')  # Оставьте это для фона
        firefox_options.add_argument('--disable-gpu')
        driver = webdriver.Firefox(options=firefox_options)

        try:
            # Открытие сайта
            driver.get('https://vsrf.ru/lk/practice/acts')

            # Ожидание полной загрузки страницы
            WebDriverWait(driver, 20).until(
                lambda driver: driver.execute_script('return document.readyState') == 'complete'
            )

            # Ожидание загрузки поля поиска
            tick_box = WebDriverWait(driver, 20).until(
                EC.presence_of_element_located((By.XPATH, '//*[@id="numberExact"]'))
            )
            if tick_box.is_selected():
                tick_box.click()

            # Ввод даты
            from_date = WebDriverWait(driver, 20).until(
                EC.presence_of_element_located((By.XPATH, '//*[@id="actDateFrom"]'))
            )
            to_date = WebDriverWait(driver, 20).until(
                EC.presence_of_element_located((By.XPATH, '//*[@id="actDateTo"]'))
            )
            from_date.send_keys(date_from)
            to_date.send_keys(date_up)

            # Ввод поискового запроса
            search_box = WebDriverWait(driver, 20).until(
                EC.presence_of_element_located((By.XPATH, '//*[@id="keywords"]'))
            )
            search_box.send_keys(search_string)

            # Нажатие на кнопку "Поиск"
            search_button = WebDriverWait(driver, 20).until(
                EC.element_to_be_clickable((By.XPATH, '//*[@id="filter-form"]/div[1]/div[1]/div[2]/div/div[1]/input'))
            )
            search_button.click()

            time.sleep(3)

            # Ожидание загрузки основного контейнера
            container = WebDriverWait(driver, 20).until(
                EC.presence_of_element_located((By.XPATH, "/html/body/div/div[2]/div/div/div[2]/div"))
            )

            # Извлечение всех элементов с классом 'vs-items-body'
            item_bodies = container.find_elements(By.XPATH, "//*[contains(@class, 'vs-items-body')]")

            table_data = []
            if item_bodies:
                for item_body in item_bodies:
                    try:
                        # Извлечение даты
                        date_element = item_body.find_element(By.XPATH, ".//*[contains(@class, 'vs-font-20')]")
                        link_names = item_body.find_elements(By.XPATH, ".//*[contains(@class, 'vs-items-label')]")
                        links = item_body.find_elements(By.TAG_NAME, 'a')

                        if links:
                            row_data = {
                                'date': date_element.text.strip(),
                                'link_name': [ln.text.strip() for ln in link_names],
                                'link_source': [link.get_attribute('href') for link in links]
                            }
                            table_data.append(row_data)
                    except NoSuchElementException:
                        print("One of the expected elements was not found.")

            return table_data

        except Exception as e:
            print(f"Error during Selenium script execution: {e}")
            traceback.print_exc()

        finally:
            driver.quit()

    def add_to_word_file(table_data: list, date_from, date_up):
        if table_data:
            doc = Document(f'Судебная_практика {date_from}-{date_up}.docx')
            bold_paragraph = doc.add_paragraph()
            bold_run = bold_paragraph.add_run('2. Верховный Суд РФ:')
            bold_run.bold = True

            for i, record in enumerate(table_data, start=1):
                paragraph = doc.add_paragraph(f'2.{i} ')
                link_text = f"Постановление от {record['date']} № {record['link_name']}"
                if record['link_source']:
                    add_hyperlink(paragraph, link_text, record['link_source'][0])

            doc.save(f'Судебная_практика {date_from}-{date_up}.docx')
            print("Документ успешно добавлен!")
        else:
            print("Не было найдено никаких решений.")

    def make_new_file(table_data: list, date_from, date_up):
        if table_data:
            doc = Document()
            doc.add_heading('Судебная практика', 0)
            bold_paragraph = doc.add_paragraph()
            bold_run = bold_paragraph.add_run('Верховный Суд РФ:')
            bold_run.bold = True

            for i, record in enumerate(table_data, start=1):
                paragraph = doc.add_paragraph(f'1.{i} ')
                link_text = f"Постановление от {record['date']} № {record['link_name']}"
                if record['link_source']:
                    add_hyperlink(paragraph, link_text, record['link_source'][0])

            doc.save(f'Судебная_практика_ВС_РФ {date_from}-{date_up}.docx')
            print("Документ успешно создан!")
        else:
            print("Не было найдено никаких решений.")

    if flag == 'add_to_file':
        add_to_word_file(parse_SC(string_to_put_in_search, date_from, date_up), date_from, date_up)
    elif flag == 'make_new_file':
        make_new_file(parse_SC(string_to_put_in_search, date_from, date_up), date_from, date_up)
