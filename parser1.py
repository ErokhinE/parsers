import sys
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.firefox.service import Service
from docx import Document
from docx.oxml import OxmlElement
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def parser1(string_to_put_in_search:str,date_from:str,date_up:str,flag:str):
    def add_hyperlink(paragraph, text, url):
        rel_id = paragraph.part.relate_to(url, 'hyperlink', is_external=True)
        hyperlink = OxmlElement('w:hyperlink', {qn('r:id'): rel_id})
        run = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = text
        run.append(t)


        run_properties = OxmlElement('w:rPr')
        # Подчеркивание
        underline = OxmlElement('w:u')
        underline.set(qn('w:val'), 'single')
        run_properties.append(underline)

        # Цвет шрифта
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')  # Синий цвет
        run_properties.append(color)

        run.append(run_properties)
        hyperlink.append(run)
        paragraph._element.append(hyperlink)

    def parse_KJRF(search_string:str,date_from:str,date_up:str)->list[list[tuple,tuple,tuple]]:
        firefox_options = webdriver.FirefoxOptions()
        firefox_options.add_argument('--headless')
        firefox_options.add_argument('--disable-gpu')
        driver = webdriver.Firefox(options=firefox_options)

        
        try:
            driver.get('https://www.ksrf.ru/ru/Decision/Pages/default.aspx')

            WebDriverWait(driver, 20).until(
                lambda driver: driver.execute_script('return document.readyState') == 'complete'
            )

            
            search_box = WebDriverWait(driver, 20).until(
                EC.presence_of_element_located((By.XPATH, '//*[@id="ctl00_m_g_8da72b0e_36c3_43d7_9458_469b90467bbc"]/div[1]/table/tbody/tr[1]/td/input[1]'))  # Change to more specific selector if needed
            )
            search_box.send_keys(search_string)

            
            from_date = WebDriverWait(driver, 20).until(
                EC.presence_of_element_located((By.XPATH, '//*[@id="ctl00_m_g_8da72b0e_36c3_43d7_9458_469b90467bbc_ctl04_Date"]'))
            )
            to_date = WebDriverWait(driver, 20).until(
                EC.presence_of_element_located((By.XPATH, '//*[@id="ctl00_m_g_8da72b0e_36c3_43d7_9458_469b90467bbc_ctl05_Date"]'))
            )
            from_date.send_keys(date_from)
            to_date.send_keys(date_up)

            
            search_button = WebDriverWait(driver, 20).until(
                EC.element_to_be_clickable((By.XPATH, '//*[@id="ctl00_m_g_8da72b0e_36c3_43d7_9458_469b90467bbc"]/div[1]/table/tbody/tr[1]/td/input[2]'))  # Using clickable condition
            )
            search_button.click()

            table = WebDriverWait(driver, 20).until(
                EC.presence_of_element_located((By.XPATH, '//*[@id="ctl00_m_g_8da72b0e_36c3_43d7_9458_469b90467bbc_gView"]'))
            )

            
            rows = table.find_elements(By.TAG_NAME, 'tr')

            
            table_data = []

        
            for row in rows:
                cells = row.find_elements(By.TAG_NAME, 'td')
                if cells:
                    row_data = []
                    for cell in cells:
                        link_element = cell.find_elements(By.TAG_NAME, 'a') 
                        if link_element:
                            link_text = link_element[0].text
                            link_url = link_element[0].get_attribute('href')
                            row_data.append((link_text, link_url))
                        else:
                            row_data.append((cell.text, None))
                    table_data.append(row_data)

            
            # for item in table_data:
            #     print(item)

            driver.quit()
            return table_data

        except Exception as e:
            print(f"Error during Selenium script execution: {e}")

        finally:
            driver.quit()



    def make_word_file_to_all_parsers(table_data:list, date_from, date_up):
        if table_data:
            
            doc = Document()
            doc.add_heading('Судебная практика', 0)
            bold_paragraph = doc.add_paragraph()
            bold_run = bold_paragraph.add_run('1. Конституционный Суд РФ:')
            bold_run.bold = True

            
            for i, case in enumerate(table_data, 1):
                case_date = case[0][0]  
                case_title = case[1][0]  
                case_number = case[2][0]  
                case_link = case[2][1]  

                
                paragraph = doc.add_paragraph(f'1.{i}. ')
                link_text = f"Постановление от {case_date} № {case_number}"
                formatted_case = f"{i}. {link_text}"

                # Добавление гиперссылки
                if case_link:  
                    add_hyperlink(paragraph, link_text, case_link)

                
                paragraph.add_run(f" {case_title}")
            doc.save(f'Судебная_практика {date_from}-{date_up}.docx')
            print("Документ успешно создан!")
        else:
            print("Не было найдено никаких решений.")

    def make_word_file(table_data:list, date_from, date_up):
        if table_data:
            
            doc = Document()
            doc.add_heading('Судебная практика', 0)
            bold_paragraph = doc.add_paragraph()
            bold_run = bold_paragraph.add_run('Конституционный Суд РФ:')
            bold_run.bold = True

            
            for i, case in enumerate(table_data, 1):
                case_date = case[0][0]  
                case_title = case[1][0]  
                case_number = case[2][0]  
                case_link = case[2][1]  

                
                paragraph = doc.add_paragraph(f'1.{i}. ')
                link_text = f"Постановление от {case_date} № {case_number}"
                formatted_case = f"{i}. {link_text}"

                # Добавление гиперссылки
                if case_link:  
                    add_hyperlink(paragraph, link_text, case_link)

                
                paragraph.add_run(f" {case_title}")

            
            doc.save(f'Судебная_практика_КС_РФ {date_from}-{date_up}.docx')
            print("Документ успешно создан!")
        else:
            print("Не было найдено никаких решений.")

    
    if flag == 'add_to_file':
        make_word_file_to_all_parsers(parse_KJRF(string_to_put_in_search,date_from,date_up), date_from, date_up)
    elif flag =='make_new_file':
        make_word_file(parse_KJRF(string_to_put_in_search,date_from,date_up), date_from, date_up)

